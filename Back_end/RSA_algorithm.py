import docx
import openpyxl
import random
from Crypto.Util import number
import math

# Thuật toán Miller-Rabin - kiểm tra số nguyên tố
def is_prime(n, k=20):
    if n <= 1:
        return False
    if n <= 3:
        return True
    if n % 2 == 0:
        return False

    # Phân tích n - 1 thành d * 2^r
    r = 0
    d = n - 1
    while d % 2 == 0:
        r += 1
        d //= 2

    # Kiểm tra k lần
    for _ in range(k):
        a = random.randint(2, n - 2)
        x = pow(a, d, n)

        if x == 1 or x == n - 1:
            continue

        for _ in range(r - 1):
            x = pow(x, 2, n)
            if x == n - 1:
                break
        else:
            return False

    return True


# Hàm sinh số nguyên tố có độ dài bit
def generate_prime(bit_length):
    while True:
        num = number.getPrime(bit_length)
        if is_prime(num):
            return num

# Hàm kiểm tra số có độ dài bit từ 512 đến 4096 và là số nguyên tố
def check_prime_with_bit_length(num):
    if num.bit_length() >= 512 and num.bit_length() <= 4096:
        if is_prime(num):
            return True
    return False

# Hàm tìm ước chung lớn nhất của a và b
def gcd(a, b):
    while b:
        a, b = b, a % b
    return a


#hàm tính nghịch đảo của a modulo
# def inverse(a, m):
#     def extended_gcd(a, b):
#         if b == 0:
#             return (a, 1, 0)
#         else:
#             d, x, y = extended_gcd(b, a % b)
#             return (d, y, x - (a // b) * y)

#     d, x, _ = extended_gcd(a, m)
#     if d == 1:
#         return x % m
#     else:
#         raise ValueError("The inverse does not exist.")


# Hàm tìm số d thỏa mãn d * e ≡ 1 (mod phi)
def find_d(e, phi):
    d = number.inverse(e, phi)
    return d % phi

# Thuật toán Euclid mở rộng - nghịch đảo modulo  
def extended_gcd(a, b):
    if b == 0:
        return a, 1, 0
    else:
        d, x, y = extended_gcd(b, a % b)
        return d, y, x - (a // b) * y


# Sinh khóa công khai và khóa bí mật
def generate_keys(bit_length):
    # Sinh số nguyên tố p và q
    p = generate_prime(bit_length // 2)
    q = generate_prime(bit_length // 2)

    n = p * q  # Tính n
    phi = (p - 1) * (q - 1)  # Tính hàm phi(n)

    # Tìm số e thỏa mãn 1 < e < phi và gcd(e, phi) = 1
    e = random.randint(1, phi - 1)
    while gcd(e, phi) != 1:
        e = random.randint(1, phi - 1)

    # Tìm số d thỏa mãn d * e ≡ 1 (mod phi)
    d = find_d(e, phi)

    public_key = (e, n)
    private_key = (d, n)

    return public_key, private_key, p, q

# hàm tính giá trị các khóa từ giá trị p và q nhập vào từ bàn phím 
def tinh_keys(p, q):
    n = p * q  # Tính n
    phi = (p - 1) * (q - 1)  # Tính hàm phi(n)

    # Tìm số e thỏa mãn 1 < e < phi và gcd(e, phi) = 1
    e = random.randint(1, phi - 1)
    while math.gcd(e, phi) != 1:
        e = random.randint(1, phi - 1)

    # Tìm số d thỏa mãn d * e ≡ 1 (mod phi)
    d = find_d(e, phi)

    public_key = (e, n)
    private_key = (d, n)

    return public_key, private_key


# Mã hoá dữ liệu
def rsa_encrypt(data, e, n):
    encrypted_data = pow(data, e, n)
    return encrypted_data

# Giải mã dữ liệu
def rsa_decrypt(encrypted_data, d, n):
    decrypted_data = pow(encrypted_data, d, n)
    return decrypted_data

# Đọc nội dung tệp tin TXT
def read_txt_file(filename):
    with open(filename, 'r') as file:
        content = file.read()
    return content

# Ghi nội dung vào tệp tin TXT
def write_txt_file(filename, content):
    with open(filename, 'w') as file:
        file.write(content)
        

# Mã hoá tệp tin TXT
def encrypt_txt_file(filename, e, n):
    content = read_txt_file(filename)
    byte_array = content.encode('utf-8')
    data = int.from_bytes(byte_array, 'big')
    encrypted_data = rsa_encrypt(data, e, n)
    write_txt_file('encrypted_' + filename, str(encrypted_data))
    return encrypted_data

# Giải mã tệp tin TXT
def decrypt_txt_file(filename, d, n):
    content = read_txt_file(filename)
    encrypted_data = int(content)
    decrypted_data = rsa_decrypt(encrypted_data, d, n)
    byte_array = decrypted_data.to_bytes((decrypted_data.bit_length() + 7) // 8, 'big')
    decrypted_content = byte_array.decode('utf-8')
    write_txt_file('decrypted_' + filename, decrypted_content)
    return decrypted_content

# Đọc nội dung tệp tin Word (docx)
def read_docx_file(filename):
    doc = docx.Document(filename)
    content = []
    for paragraph in doc.paragraphs:
        content.append(paragraph.text)
    return '\n'.join(content)

# Ghi nội dung vào tệp tin Word (docx)
def write_docx_file(filename, content):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(filename)

# Mã hoá tệp tin Word (docx)
def encrypt_docx_file(filename, e, n):
    content = read_docx_file(filename)
    byte_array = content.encode('utf-8')
    data = int.from_bytes(byte_array, 'big')
    encrypted_data = rsa_encrypt(data, e, n)
    write_docx_file('encrypted_' + filename, str(encrypted_data))
    return encrypted_data

# Giải mã tệp tin Word (docx)
def decrypt_docx_file(filename, d, n):
    content = read_docx_file(filename)
    encrypted_data = int(content)
    decrypted_data = rsa_decrypt(encrypted_data, d, n)
    byte_array = decrypted_data.to_bytes((decrypted_data.bit_length() + 7) // 8, 'big')
    decrypted_content = byte_array.decode('utf-8', errors='ignore')
    write_docx_file('decrypted_' + filename, decrypted_content)
    return decrypted_content



# Đọc nội dung tệp tin Excel (XLSX)
def read_xlsx_file(filename):
    workbook = openpyxl.load_workbook(filename)
    content = []
    for sheet in workbook.sheetnames:
        worksheet = workbook[sheet]
        for row in worksheet.iter_rows():
            for cell in row:
                content.append(str(cell.value))
    return '\n'.join(content)

# Ghi nội dung vào tệp tin Excel (XLSX)
def write_xlsx_file(filename, content):
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    lines = content.split('\n')
    for i, line in enumerate(lines, start=1):
        worksheet.cell(row=i, column=1, value=line)
    workbook.save(filename)

# Mã hoá tệp tin Excel (XLSX)
def encrypt_xlsx_file(filename, e, n):
    content = read_xlsx_file(filename)
    byte_array = content.encode('utf-8')
    data = int.from_bytes(byte_array, 'big')
    encrypted_data = rsa_encrypt(data, e, n)
    write_xlsx_file('encrypted_' + filename, str(encrypted_data))
    return encrypted_data

# Giải mã tệp tin Excel (XLSX)
def decrypt_xlsx_file(filename, d, n):
    content = read_xlsx_file(filename)
    encrypted_data = int(content)
    decrypted_data = rsa_decrypt(encrypted_data, d, n)
    byte_array = decrypted_data.to_bytes((decrypted_data.bit_length() + 7) // 8, 'big')
    decrypted_content = byte_array.decode('utf-8')
    write_xlsx_file('decrypted_' + filename, decrypted_content)
    return decrypted_content


# # Sinh khóa công khai và khóa bí mật
# bit_length = 2048
# (public_key, private_key, p, q) = generate_keys(bit_length)

# Lưu khóa công khai và khóa bí mật vào tệp tin
def save_file(public_key, private_key, p, q):
    with open('public_key.txt', 'w') as file:
        file.write(f"e: {public_key[0]}\n")
        file.write(f"n: {public_key[1]}")

    with open('private_key.txt', 'w') as file:
        file.write(f"d: {private_key[0]}\n")
        file.write(f"n: {private_key[1]}\n")
        file.write(f"p: {p}\n")
        file.write(f"q: {q}")

# Đọc khóa công khai và khóa bí mật từ tệp tin
def read_file(file_public,file_private):
    with open(file_public, 'r') as file:
        lines = file.readlines()
        e = int(lines[0].split(':')[1].strip())
        n = int(lines[1].split(':')[1].strip())

    with open(file_private, 'r') as file:
        lines = file.readlines()
        d = int(lines[0].split(':')[1].strip())
        n = int(lines[1].split(':')[1].strip())
        p = int(lines[2].split(':')[1].strip())
        q = int(lines[3].split(':')[1].strip())

# Mã hoá và giải mã các tệp tin

# Mã hoá tệp tin
def encrypt_file(filename, e, n):
    content = ""
    file_extension = filename.split('.')[-1]
    if file_extension == 'txt':
        content = str(encrypt_txt_file(filename, e, n))
    elif file_extension == 'docx':     
        content = str(encrypt_docx_file(filename, e, n))
    elif file_extension == 'xlsx':
        content = str(encrypt_xlsx_file(filename, e, n))
        
    else:
        print(f"Định dạng tệp không̣ được hỗ trợ")
    return content

# Giải mã tệp tin
def decrypt_file(filename, d, n):
    resutl = ''
    file_extension = filename.split('.')[-1]
    if file_extension == 'txt':
        result = str(decrypt_txt_file(filename, d, n))
    elif file_extension == 'docx':
        result = str(decrypt_docx_file(filename, d, n))   
    elif file_extension == 'xlsx':
        result = str(decrypt_xlsx_file(filename, d, n))    
    else:
        print(f"Định dạng tệp không̣ được hỗ trợ")
    return result

